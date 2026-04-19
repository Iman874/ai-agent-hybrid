import io
import re
import logging

import markdown
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.utils.errors import ExportError

logger = logging.getLogger("ai-agent-hybrid.export")

SUPPORTED_FORMATS = {"docx", "pdf", "md"}


class DocumentExporterService:
    """Konversi Markdown TOR ke berbagai format dokumen.

    Service ini stateless — tidak memiliki dependency injection.
    Semua method menerima string Markdown dan mengembalikan bytes.
    """

    def export(self, md_text: str, fmt: str) -> bytes:
        """Dispatcher utama. Routing ke method spesifik berdasarkan format.

        Args:
            md_text: String konten TOR dalam format Markdown.
            fmt: Format target — "docx", "pdf", atau "md".

        Returns:
            bytes: File content sebagai binary.

        Raises:
            ExportError: Jika format tidak didukung atau konversi gagal.
        """
        fmt = fmt.lower().strip()
        if fmt not in SUPPORTED_FORMATS:
            raise ExportError(
                message=f"Format '{fmt}' tidak didukung.",
                details=f"Format yang didukung: {', '.join(sorted(SUPPORTED_FORMATS))}",
            )
        method_map = {
            "md": self.export_to_md,
            "pdf": self.export_to_pdf,
            "docx": self.export_to_docx,
        }
        try:
            result = method_map[fmt](md_text)
            logger.info(f"Export sukses: format={fmt}, size={len(result)} bytes")
            return result
        except ExportError:
            raise
        except Exception as e:
            logger.error(f"Export gagal: format={fmt}, error={e}")
            raise ExportError(
                message=f"Gagal mengekspor ke format {fmt}.",
                details=str(e),
            )

    def export_to_md(self, md_text: str) -> bytes:
        """Return raw Markdown text sebagai UTF-8 bytes. Tanpa konversi."""
        return md_text.encode("utf-8")

    def export_to_pdf(self, md_text: str) -> bytes:
        """Konversi Markdown → HTML → PDF via xhtml2pdf.

        Logika dimigrasikan dari streamlit_app/utils/formatters.py
        dengan styling yang dipertahankan.
        """
        html_body = markdown.markdown(
            md_text, extensions=["tables", "fenced_code"]
        )
        styled_html = f"""<html><head><style>
            body {{ font-family: 'Helvetica', Arial, sans-serif;
                   font-size: 12pt; line-height: 1.6; color: #222;
                   margin: 40px; }}
            h1 {{ font-size: 18pt; text-align: center; margin-bottom: 20px; }}
            h2 {{ font-size: 14pt; border-bottom: 1px solid #ccc;
                  padding-bottom: 5px; margin-top: 25px; }}
            h3 {{ font-size: 12pt; margin-top: 20px; }}
            table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
            th {{ background: #f4f4f4; font-weight: bold; }}
            ul, ol {{ margin-left: 20px; }}
        </style></head><body>{html_body}</body></html>"""

        buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(io.StringIO(styled_html), dest=buffer)
        if pisa_status.err:
            raise ExportError(
                message="Gagal membuat PDF.",
                details=f"xhtml2pdf error count: {pisa_status.err}",
            )
        return buffer.getvalue()

    def export_to_docx(self, md_text: str) -> bytes:
        """Konversi Markdown → python-docx Document.

        Parsing regex sederhana yang menangani elemen umum TOR:
        headings, bold, italic, bullet lists, ordered lists, tables.
        """
        doc = Document()

        # --- Page style ---
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)

        lines = md_text.split("\n")
        i = 0
        table_buffer: list[str] = []

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # --- Flush table jika baris bukan tabel ---
            if table_buffer and not stripped.startswith("|"):
                self._flush_table(doc, table_buffer)
                table_buffer = []

            # --- Heading ---
            heading_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                doc.add_heading(text, level=min(level, 4))
                i += 1
                continue

            # --- Table row ---
            if stripped.startswith("|") and stripped.endswith("|"):
                # Skip separator rows (|---|---|)
                if re.match(r"^\|[\s\-:|]+\|$", stripped):
                    i += 1
                    continue
                table_buffer.append(stripped)
                i += 1
                continue

            # --- Unordered list ---
            list_match = re.match(r"^[-*]\s+(.+)$", stripped)
            if list_match:
                text = list_match.group(1)
                p = doc.add_paragraph(style="List Bullet")
                self._apply_inline_formatting(p, text)
                i += 1
                continue

            # --- Ordered list ---
            olist_match = re.match(r"^\d+\.\s+(.+)$", stripped)
            if olist_match:
                text = olist_match.group(1)
                p = doc.add_paragraph(style="List Number")
                self._apply_inline_formatting(p, text)
                i += 1
                continue

            # --- Horizontal rule ---
            if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", stripped):
                # Add a thin horizontal-like paragraph
                p = doc.add_paragraph("_" * 50)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                i += 1
                continue

            # --- Empty line (skip) ---
            if not stripped:
                i += 1
                continue

            # --- Regular paragraph ---
            p = doc.add_paragraph()
            self._apply_inline_formatting(p, stripped)
            i += 1

        # Flush remaining table
        if table_buffer:
            self._flush_table(doc, table_buffer)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    # --- Private helpers ---

    def _flush_table(self, doc: Document, rows: list[str]) -> None:
        """Parse baris-baris tabel Markdown menjadi tabel Word."""
        parsed: list[list[str]] = []
        for row in rows:
            cells = [c.strip() for c in row.strip("|").split("|")]
            parsed.append(cells)

        if not parsed:
            return

        n_cols = max(len(r) for r in parsed)
        table = doc.add_table(rows=len(parsed), cols=n_cols)
        table.style = "Table Grid"

        for r_idx, row_data in enumerate(parsed):
            for c_idx, cell_text in enumerate(row_data):
                if c_idx < n_cols:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_text
                    # Bold header row
                    if r_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

    def _apply_inline_formatting(self, paragraph, text: str) -> None:
        """Parse inline Markdown (bold, italic) dan terapkan ke paragraph runs."""
        # Pattern: **bold**, *italic*
        parts = re.split(r"(\*{1,2}[^*]+\*{1,2})", text)
        for part in parts:
            if not part:
                continue
            bold_match = re.match(r"^\*\*(.+?)\*\*$", part)
            italic_match = re.match(r"^\*([^*]+)\*$", part)
            if bold_match:
                run = paragraph.add_run(bold_match.group(1))
                run.bold = True
            elif italic_match:
                run = paragraph.add_run(italic_match.group(1))
                run.italic = True
            else:
                paragraph.add_run(part)
