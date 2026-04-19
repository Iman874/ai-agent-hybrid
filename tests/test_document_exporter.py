"""
Unit tests untuk DocumentExporterService.
Tidak memerlukan database atau server.
"""
import io
import pytest
from docx import Document as DocxDocument

from app.services.document_exporter import DocumentExporterService
from app.utils.errors import ExportError


# --- Fixture ---

@pytest.fixture
def exporter():
    """Instance DocumentExporterService untuk testing."""
    return DocumentExporterService()


@pytest.fixture
def sample_markdown():
    """Sample TOR Markdown content yang realistis."""
    return """# TERM OF REFERENCE (TOR)
# Workshop Penerapan AI untuk ASN

## 1. Latar Belakang

Perkembangan teknologi **Artificial Intelligence** (AI) telah mengubah berbagai sektor.
Pemerintah perlu mempersiapkan ASN yang *kompeten* dalam pemanfaatan AI.

## 2. Tujuan

- Meningkatkan pemahaman ASN tentang AI
- Membangun kapasitas digital aparatur
- Mendorong inovasi layanan publik

## 3. Ruang Lingkup

| No | Komponen | Keterangan |
|----|----------|------------|
| 1  | Pelatihan | 3 hari intensif |
| 2  | Workshop | Hands-on project |
| 3  | Evaluasi | Pre & post test |

## 4. Timeline

1. Persiapan: Minggu 1-2
2. Pelaksanaan: Minggu 3
3. Evaluasi: Minggu 4

---

## 5. Penutup

Demikian TOR ini disusun sebagai acuan pelaksanaan kegiatan.
"""


# --- Tests: export_to_md ---

class TestExportMd:
    def test_returns_bytes(self, exporter, sample_markdown):
        result = exporter.export_to_md(sample_markdown)
        assert isinstance(result, bytes)

    def test_content_identical(self, exporter, sample_markdown):
        """MD export harus identik dengan input setelah decode."""
        result = exporter.export_to_md(sample_markdown)
        assert result.decode("utf-8") == sample_markdown

    def test_utf8_encoding(self, exporter):
        """Harus handle karakter Indonesia dengan benar."""
        text = "# TOR Pengadaan Barang — Rp 50.000.000"
        result = exporter.export_to_md(text)
        assert "Rp 50.000.000" in result.decode("utf-8")

    def test_empty_input(self, exporter):
        result = exporter.export_to_md("")
        assert result == b""


# --- Tests: export_to_pdf ---

class TestExportPdf:
    def test_returns_bytes(self, exporter, sample_markdown):
        result = exporter.export_to_pdf(sample_markdown)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_pdf_magic_bytes(self, exporter, sample_markdown):
        """PDF file harus dimulai dengan magic bytes %PDF."""
        result = exporter.export_to_pdf(sample_markdown)
        assert result[:4] == b"%PDF", f"Got: {result[:20]}"

    def test_minimal_input(self, exporter):
        result = exporter.export_to_pdf("# Hello")
        assert result[:4] == b"%PDF"


# --- Tests: export_to_docx ---

class TestExportDocx:
    def test_returns_bytes(self, exporter, sample_markdown):
        result = exporter.export_to_docx(sample_markdown)
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_valid_docx_file(self, exporter, sample_markdown):
        """Output harus bisa di-load sebagai valid DOCX."""
        result = exporter.export_to_docx(sample_markdown)
        doc = DocxDocument(io.BytesIO(result))
        assert len(doc.paragraphs) > 0 or len(doc.tables) > 0

    def test_heading_preserved(self, exporter):
        """Heading Markdown harus menjadi heading di DOCX."""
        result = exporter.export_to_docx("# Judul Utama\n\n## Sub Judul")
        doc = DocxDocument(io.BytesIO(result))
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert "Judul Utama" in headings
        assert "Sub Judul" in headings

    def test_table_created(self, exporter):
        """Tabel Markdown harus menjadi tabel di DOCX."""
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        result = exporter.export_to_docx(md)
        doc = DocxDocument(io.BytesIO(result))
        assert len(doc.tables) >= 1

    def test_bold_formatting(self, exporter):
        """**Bold** Markdown harus menjadi bold run di DOCX."""
        result = exporter.export_to_docx("Ini **tebal** sekali")
        doc = DocxDocument(io.BytesIO(result))
        bold_runs = []
        for p in doc.paragraphs:
            for run in p.runs:
                if run.bold:
                    bold_runs.append(run.text)
        assert "tebal" in bold_runs

    def test_bullet_list(self, exporter):
        """List Markdown harus menjadi list style di DOCX."""
        md = "- Item satu\n- Item dua\n- Item tiga"
        result = exporter.export_to_docx(md)
        doc = DocxDocument(io.BytesIO(result))
        list_items = [p.text for p in doc.paragraphs if "List" in (p.style.name or "")]
        assert len(list_items) >= 3


# --- Tests: export() dispatcher ---

class TestExportDispatcher:
    def test_md_dispatch(self, exporter):
        result = exporter.export("# Test", "md")
        assert result == b"# Test"

    def test_pdf_dispatch(self, exporter):
        result = exporter.export("# Test", "pdf")
        assert result[:4] == b"%PDF"

    def test_docx_dispatch(self, exporter):
        result = exporter.export("# Test", "docx")
        doc = DocxDocument(io.BytesIO(result))
        assert doc is not None

    def test_invalid_format_raises(self, exporter):
        with pytest.raises(ExportError) as exc_info:
            exporter.export("# Test", "xlsx")
        assert "tidak didukung" in exc_info.value.message

    def test_case_insensitive(self, exporter):
        """Format harus case-insensitive."""
        result = exporter.export("# Test", "PDF")
        assert result[:4] == b"%PDF"

    def test_whitespace_trimmed(self, exporter):
        """Format dengan whitespace harus di-trim."""
        result = exporter.export("# Test", " md ")
        assert result == b"# Test"
